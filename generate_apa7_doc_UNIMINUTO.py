import os
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_double_spacing(paragraph):
    paragraph.paragraph_format.line_spacing = 2.0

def create_apa7_uniminuto_document(content_file, output_path, image_path=None, bibliography_file=None):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Read content
    try:
        with open(content_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: No se pudo encontrar el archivo de contenido en: {content_file}")
        return
    except Exception as e:
        print(f"Error al leer el archivo de contenido: {e}")
        return

    # Cover Page
    title_line_1 = "Diagrama de Flujo"
    title_line_2 = "Procesos de Entrada y Salida de Datos en el Computador"
    
    doc.add_paragraph().add_run("\n\n\n\n\n")
    titulo1 = doc.add_paragraph(title_line_1)
    titulo1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo1.runs[0].bold = True
    set_double_spacing(titulo1)
    
    titulo2 = doc.add_paragraph(title_line_2)
    titulo2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo2.runs[0].bold = True
    set_double_spacing(titulo2)
    
    doc.add_paragraph().add_run("\n\n\n\n") # 4 enters

    autor = doc.add_paragraph("Gerardo Castillo Martinez")
    autor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_double_spacing(autor)
    
    uni = doc.add_paragraph("Universidad Minuto de Dios")
    uni.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_double_spacing(uni)

    prog = doc.add_paragraph("Introducción a la Ingeniería de Software")
    prog.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_double_spacing(prog)

    prof = doc.add_paragraph("James Gabriel Jaramillo Zambrano")
    prof.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_double_spacing(prof)

    from datetime import date
    today = date.today()
    fecha_actual = today.strftime("%d de %B de %Y").replace("January", "Enero").replace("February", "Febrero").replace("March", "Marzo").replace("April", "Abril").replace("May", "Mayo").replace("June", "Junio").replace("July", "Julio").replace("August", "Agosto").replace("September", "Septiembre").replace("October", "Octubre").replace("November", "Noviembre").replace("December", "Diciembre")
    fecha = doc.add_paragraph(fecha_actual)
    fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_double_spacing(fecha)

    doc.add_page_break()

    # Process Body
    body_lines = [line.strip() for line in lines if line.strip()]
    
    # Find sections
    intro_start = -1
    dev_start = -1
    conc_start = -1
    
    for i, line in enumerate(body_lines):
        if line == "## Introducción":
            intro_start = i
        elif line == "## Desarrollo":
            dev_start = i
        elif line == "## Conclusiones":
            conc_start = i

    # --- Introduction ---
    if intro_start != -1:
        intro_title = doc.add_paragraph("Procesos de Entrada y Salida de Datos en el Computador")
        intro_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        intro_title.runs[0].bold = True
        set_double_spacing(intro_title)
        doc.add_paragraph()
        
        intro_end = dev_start if dev_start != -1 else conc_start if conc_start != -1 else len(body_lines)
        for i in range(intro_start + 1, intro_end):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            set_double_spacing(p)
            p.add_run(body_lines[i])

    # --- Development ---
    if dev_start != -1:
        doc.add_page_break()
        dev_title = doc.add_paragraph("Desarrollo")
        dev_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        dev_title.runs[0].bold = True
        set_double_spacing(dev_title)
        doc.add_paragraph()
        
        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
            
        dev_end = conc_start if conc_start != -1 else len(body_lines)
        for i in range(dev_start + 1, dev_end):
            line = body_lines[i]
            if line.startswith('### '):
                title = line.replace('### ', '').strip()
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                runner = p.add_run(title)
                runner.bold = True
                set_double_spacing(p)
            elif line.startswith('#### '):
                title = line.replace('#### ', '').strip()
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                runner = p.add_run(title)
                runner.bold = True
                runner.italic = True
                set_double_spacing(p)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                set_double_spacing(p)
                p.add_run(line.replace('**', ''))

    # --- Conclusions ---
    if conc_start != -1:
        doc.add_page_break()
        conc_title = doc.add_paragraph("Conclusiones")
        conc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        conc_title.runs[0].bold = True
        set_double_spacing(conc_title)
        doc.add_paragraph()
        
        for i in range(conc_start + 1, len(body_lines)):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            set_double_spacing(p)
            p.add_run(body_lines[i])

    # Add Bibliography if provided
    if bibliography_file and os.path.exists(bibliography_file):
        doc.add_page_break()
        references_title = doc.add_paragraph("Referencias")
        references_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        references_title.runs[0].bold = True
        set_double_spacing(references_title)
        
        with open(bibliography_file, 'r', encoding='utf-8') as f:
            bibliography_content = f.read()
            # Split by <div class="csl-entry"> to process each entry
            entries = bibliography_content.split('<div class="csl-entry">')
            for entry in entries:
                if entry.strip():
                    # Remove </div> and any leading/trailing whitespace
                    clean_entry = entry.replace('</div>', '').strip()
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Inches(0.5) # Hanging indent
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.line_spacing = 2.0 # Double spacing
                    p.add_run(clean_entry)

    # Save document
    try:
        doc.save(output_path)
        print(f"Documento final generado con éxito en: {output_path}")
    except Exception as e:
        print(f"Error al guardar el documento: {e}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Genera un documento .docx en formato APA 7 Uniminuto a partir de un archivo de texto.')
    parser.add_argument('--content_file', required=True, help='Ruta al archivo .txt con el contenido del ensayo.')
    parser.add_argument('--output_path', required=True, help='Ruta donde se guardará el archivo .docx final.')
    parser.add_argument('--image_path', required=False, help='Ruta a la imagen del diagrama de flujo.')
    parser.add_argument('--bibliography_file', required=True, help='Ruta al archivo .txt con la bibliografía formateada en APA 7.')
    
    args = parser.parse_args()
    
    create_apa7_uniminuto_document(args.content_file, args.output_path, args.image_path, args.bibliography_file)